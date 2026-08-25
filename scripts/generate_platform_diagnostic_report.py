from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "reports/mlb-analyst-platform-diagnostic-2026-08-25.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], head, True)
        shade(table.rows[0].cells[i], "1F4E78")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(6)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Aptos Display"
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 120)
styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 120)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MLB Analyst Platform\nOperational Diagnostic & Remediation Report")
run.bold = True
run.font.name = "Aptos Display"
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(31, 78, 120)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Prepared from live application behavior, database evidence, source-health responses, and runtime logs\n"
                 "Assessment date: August 25, 2026 (America/New_York)").italic = True
doc.add_paragraph()

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The platform is partially functional but is not currently operating as a smooth, publish-ready daily workflow. "
    "Its strongest behavior is safety: it now detects incomplete Ballpark Pal coverage and prevents a research freeze "
    "or operational publication instead of silently publishing incomplete market ranks. That is the correct safety "
    "behavior. However, the user experience is still poor because the dashboard is blocked, weather visibly appears "
    "missing despite being stored, official game metadata is stale and incomplete, and provider/player reconciliation "
    "does not yet cover all projected hitters."
)
doc.add_paragraph(
    "There are two different meanings of “working” in the current environment. The API server starts, the five market "
    "engines can execute, the market board can materialize, and Round Robin logic has data to read. In contrast, the "
    "operational orchestration layer correctly refuses to declare the slate ready. The distinction matters: a manually "
    "generated board is not the same thing as a verified, complete, auditable daily slate. The platform is presently in "
    "the latter state: usable for diagnosis and research inspection, but not ready for routine operational use."
)

add_table(doc,
          ["Area", "Current state", "Operational meaning"],
          [
              ["API service", "Running after restart", "Service can answer requests, but process restarts have appeared in logs."],
              ["Daily readiness", "BLOCKED", "Latest workflow is cancelled; the slate is not eligible for a controlled freeze/publish."],
              ["MLB official source", "STALE", "Latest effective date is Aug. 23 while the investigated slate is Aug. 25."],
              ["FantasyPros projected slate", "FRESH", "Projected lineups and embedded weather were ingested for the date."],
              ["Ballpark Pal research", "PARTIAL", "Coverage receipt finds 35 projected-hitter gaps and one park gap."],
              ["Weather storage", "Present", "25 FantasyPros observations exist across all 15 games."],
              ["Weather dashboard display", "Broken", "All 15 cards return weather and roof as “NOT FOUND.”"],
          ],
          [1.4, 1.6, 3.5])

doc.add_heading("2. Current State and Evidence", level=1)
doc.add_heading("2.1 Slate readiness", level=2)
doc.add_paragraph(
    "The live Today and Data Health responses report the slate as BLOCKED. The primary reason is explicit: "
    "“The latest workflow is cancelled.” The response also notes that optional enrichment is incomplete. This means "
    "the application is not merely showing a cosmetic warning; its orchestration state is intentionally preventing a "
    "normal operational conclusion."
)
doc.add_paragraph(
    "A cancelled workflow should be treated as an unresolved operational incident until its cause is understood and a "
    "new complete run succeeds. Manually refreshing individual engines after a cancelled orchestration can produce "
    "inspectable research rows, but cannot replace the orchestration record as the official readiness authority."
)

doc.add_heading("2.2 Research-source coverage", level=2)
doc.add_paragraph(
    "The latest Ballpark Pal run received 315 source rows, normalized 314, and rejected one. More importantly, the new "
    "coverage receipt reports zero missing games, zero missing starters, 35 missing projected hitters, and one missing "
    "park record. The run is therefore PARTIAL. This is an accurate and desirable classification: no provider row "
    "should be assumed to apply to a player merely because the provider returned a similar number of total rows."
)
doc.add_paragraph(
    "The current data proves that this is not just a simple count issue. FantasyPros has 270 distinct projected hitters "
    "for the slate and Ballpark Pal has 270 distinct hitter snapshots, but only 235 canonical player IDs overlap. "
    "Thirty-five projected hitters have no Ballpark Pal snapshot with the same canonical MLB player ID. This signals an "
    "identity/reconciliation mismatch, genuine provider omission, or a combination of both. It must be resolved per "
    "player and per game; simply accepting the matching total count would hide the defect."
)

doc.add_heading("2.3 Weather evidence", level=2)
doc.add_paragraph(
    "Weather is not absent from the database. The investigation found 25 FantasyPros weather-observation records "
    "covering all 15 games on the slate. Every one of those 25 records includes both a temperature and wind-speed "
    "value. The additional weather refresh correctly reports SUCCESS with zero newly written observations because "
    "FantasyPros observations already exist and are the preferred pregame source. The refresh result is confusing, but "
    "it is not evidence that no weather data exists."
)
doc.add_paragraph(
    "The actual display defect is in the Today endpoint. Its game mapping sets the roof field to the literal string "
    "“NOT FOUND” and sets the weather field to the literal string “NOT FOUND”; it does not join or query "
    "game_weather_observations. As a result, all 15 game cards display missing weather even though the stored data is "
    "available. This is a response-contract/UI integration bug, not an upstream weather-data absence."
)

doc.add_heading("3. Confirmed Bugs and Defects", level=1)
add_table(doc,
          ["ID", "Severity", "Confirmed issue", "Evidence and impact"],
          [
              ["B-01", "Critical", "Daily slate remains operationally BLOCKED after a cancelled workflow.", "Both Today and Data Health state that the latest workflow is cancelled. A cancelled run cannot be treated as an approved slate readiness record."],
              ["B-02", "Critical", "Ballpark Pal coverage does not include every projected hitter or park context.", "Coverage receipt: 35 hitter gaps and 1 park gap. Publishing would create incomplete or empty evidence paths for affected candidates."],
              ["B-03", "High", "Weather and roof are hard-coded as “NOT FOUND” in the Today response.", "15 games display missing weather/roof while 25 FantasyPros weather observations cover all 15 games in storage."],
              ["B-04", "High", "Official MLB schedule metadata is stale and has no venue or start-time values for the active slate.", "MLB source health is stale (effective date Aug. 23); all 15 Aug. 25 games have null venue_id and null start_time_utc. This disables Open-Meteo venue forecasts."],
              ["B-05", "High", "Provider-player identity reconciliation is incomplete.", "270 projected hitters and 270 provider hitters produce only 235 canonical-ID matches. Counts alone are misleading."],
              ["B-06", "Medium", "Weather refresh status is not sufficiently explanatory.", "It reports SUCCESS and zero observations written because it relies on existing FantasyPros fallback observations, but does not surface fallback coverage in the compact health summary."],
              ["B-07", "Medium", "Runtime logs show database connection termination events that crash an active client path.", "Deployment logs include “terminating connection due to administrator command” followed by an emitted client error stack. This needs resilience review, even if the restart initiated the termination."],
              ["B-08", "Medium", "Manual engine materialization can be mistaken for an operationally approved slate.", "Five engines and the board can write research rows while orchestration is blocked. The UI and runbook need clearer separation between exploratory output and publishable readiness."],
              ["B-09", "Low", "Weather-source precedence is inconsistent between single-game and slate retrieval.", "Single-game weather explicitly prioritizes FantasyPros; slate retrieval uses the newest record regardless of source. This can create inconsistent display/scoring provenance after Open-Meteo retries."],
          ],
          [0.55, 0.65, 2.25, 3.05])

doc.add_heading("4. Detailed Weather Diagnosis", level=1)
doc.add_heading("4.1 What is working", level=2)
add_bullets(doc, [
    "FantasyPros projected-lineup ingestion captures weather text, temperature, wind speed, wind direction, and rain chance when supplied.",
    "The database has 25 FantasyPros observations spanning every one of the 15 scheduled games. Multiple observations are expected because weather observations are append-only for auditability.",
    "The market engines can read slate weather through the weather accessor and apply bounded, market-specific adjustments. Closed roofs are modeled as neutral rather than incorrectly treated as absent.",
    "The supplemental Open-Meteo refresh correctly avoids cloning existing FantasyPros observations into a second source row."
])
doc.add_heading("4.2 Why weather is not showing in the site", level=2)
doc.add_paragraph(
    "The Today endpoint is the API response powering the game cards. It selects game and lineup fields but does not select "
    "weather observations. When constructing the response it assigns both roof and weather to constant “NOT FOUND” "
    "strings. The frontend is therefore faithfully rendering a broken response contract. No amount of clicking the "
    "Open-Meteo retry button will fix this particular symptom, because the missing step is not ingestion; it is retrieval "
    "and formatting for the Today endpoint."
)
doc.add_heading("4.3 Why supplemental Open-Meteo weather is also constrained", level=2)
doc.add_paragraph(
    "All 15 active game records have null venue IDs and null start times. Open-Meteo requires latitude and longitude, "
    "and the application’s venue hydration uses the official venue ID. Without an official venue, the application cannot "
    "safely request a venue-specific forecast. The system falls back to existing FantasyPros weather when available, "
    "which is sensible, but it cannot improve weather coverage independently until official schedule/venue metadata is "
    "fresh and reconciled."
)
doc.add_heading("4.4 Weather fixes required", level=2)
add_bullets(doc, [
    "Modify the Today query to join the preferred weather observation per game, then return a formatted weather summary, roof state, source, retrieval time, and an explicit “unavailable” reason only when no observation exists.",
    "Use a consistent preference rule in both getGameWeather and getSlateWeather: FantasyPros pregame observation first when available, then the newest Open-Meteo observation, while retaining both for audit history.",
    "Refresh the official MLB schedule before projected-lineup processing and persist canonical venue ID, start time, and roof metadata for every game.",
    "Expose the weather fallback status in Data Health: for example, “15/15 FantasyPros weather records available; 0 Open-Meteo fetches required because official venue geometry is unavailable.”",
    "Add API and UI tests showing (a) a FantasyPros weather record appears on a game card, (b) a dome displays neutral/closed rather than missing, and (c) a genuinely missing record renders a truthful unavailable state."
])

doc.add_heading("5. Root-Cause Analysis for Smooth Daily Operation", level=1)
doc.add_heading("5.1 The platform lacks a complete canonical game spine", level=2)
doc.add_paragraph(
    "The central operational issue is the gap between three representations of a game: official MLB schedule data, "
    "FantasyPros projected-slate data, and Ballpark Pal daily research data. The application has valid game identifiers "
    "for its projected slate, but the active official records do not contain venue or start-time data and Ballpark Pal "
    "does not reconcile to all projected player IDs. A daily analyst system cannot rely on aggregate counts; it needs a "
    "canonical game and player identity spine that all sources reference consistently."
)
doc.add_heading("5.2 Safety has improved, but completion has not", level=2)
doc.add_paragraph(
    "The recent coverage gate is a positive correction. It identifies omissions and marks the research run PARTIAL, "
    "rather than treating nonzero source rows as success. However, the system still needs the remediation path that "
    "turns partial coverage into complete coverage. Until then, it will safely refuse to publish—correct behavior, but "
    "not a smooth daily operating experience."
)
doc.add_heading("5.3 Provider IDs cannot be assumed to be MLB IDs", level=2)
doc.add_paragraph(
    "The 35-player mismatch shows why direct numeric identity assumptions are not sufficient. A resolution layer should "
    "first use a provider-supplied canonical identifier when documented, then use an exact normalized name plus team plus "
    "game match, and only accept a mapping when it is unique. Ambiguous cases—especially common names, midseason team "
    "changes, and doubleheaders—must remain quarantined rather than guessed. The mapping must be stored with provider "
    "provenance and an audit reason."
)

doc.add_heading("6. Prioritized Remediation Plan", level=1)
add_table(doc,
          ["Priority", "Action", "Owner outcome / acceptance criteria"],
          [
              ["P0", "Restore and verify the canonical official MLB game feed.", "A fresh current-date run fills venue_id and start_time_utc for every scheduled game; no stale official-source warning remains."],
              ["P0", "Fix Today weather/roof retrieval.", "Every game with a stored observation shows source-backed weather and roof values. “NOT FOUND” is used only when the database has no usable observation."],
              ["P0", "Resolve the 35 player and one park coverage gaps.", "Coverage receipt reports 0 game, 0 hitter, 0 pitcher, and 0 park gaps before the run can be SUCCESS."],
              ["P1", "Build auditable provider identity reconciliation.", "Unambiguous Ballpark Pal-to-MLB mappings are persisted with evidence; ambiguous mappings are quarantined and displayed in Data Health."],
              ["P1", "Reconcile doubleheaders using unique game/time context.", "No provider game can be assigned across a same-team doubleheader without a unique canonical match."],
              ["P1", "Make freeze state visible and unambiguous in UI.", "The dashboard distinguishes manual/research-only rows from a completed orchestration run that is approved to freeze/publish."],
              ["P2", "Harden database connection handling.", "Connection termination does not generate unhandled client-error crashes; retry/reconnect behavior and shutdown handling are logged and tested."],
              ["P2", "Improve weather health wording.", "The UI reports source coverage and fallback use, not only new observations written by the optional refresh."],
              ["P2", "Add fixture-driven regression tests.", "Tests cover missing game, hitter, pitcher, park, venue, weather, ambiguous doubleheader, and successful complete-slate cases."],
          ],
          [0.6, 2.25, 3.65])

doc.add_heading("7. Recommended Execution Sequence", level=1)
for i, text in enumerate([
    "Run a controlled MLB official refresh for the active slate date and inspect all 15 games for venue, start time, and canonical teams. Do not proceed if any record remains null.",
    "Repair the Today API response so it reads the preferred stored weather observation and returns weather/roof/source data. Confirm the UI displays it on all game cards.",
    "Create a provider identity reconciliation report for the 35 unmatched projected hitters. For each, classify the cause as provider absence, stable mapping needed, ambiguous mapping, or incorrect projected lineup identity.",
    "Create and verify a park-context reconciliation report for the one missing game. Confirm the provider game maps to the correct official game and applicable park multipliers.",
    "Rerun Ballpark Pal ingestion. Require the coverage receipt to be zero across game, hitter, starter, and park categories before accepting SUCCESS.",
    "Rerun all five engines, market-board materialization, and RR1–RR5 only after the research run is complete.",
    "Start a fresh orchestration run and allow it to become the readiness authority. Confirm Data Health reports READY rather than simply showing manually materialized rows.",
    "Review error logs after the run. If database connection terminations recur without an intentional deployment/restart, treat them as a production stability incident."
], 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    p.add_run(text)

doc.add_heading("8. Acceptance Checklist: What “Running Smoothly” Means", level=1)
add_bullets(doc, [
    "The API, web artifact, and background workflow are running with no recurring unhandled runtime errors.",
    "Official MLB data is current for the slate date, including game time and venue for every game.",
    "FantasyPros projected lineups are current, and their identities resolve to canonical players without blocking issues.",
    "Ballpark Pal is SUCCESS, not PARTIAL, and reports complete game/hitter/pitcher/park coverage.",
    "The dashboard visibly displays weather, roof context, source, and freshness for every game with an observation.",
    "All five market engines run after a complete research ingest and no candidate gets empty provider evidence due to a hidden mapping failure.",
    "The board and Round Robin views identify their run/readiness state and are not confused with a manually refreshed but unapproved slate.",
    "The latest orchestration run completes its health check and becomes the persisted readiness authority.",
    "Regression tests reproduce the previously observed missing-weather and incomplete-coverage failures and prove the system blocks safely."
])

doc.add_heading("9. Final Assessment", level=1)
doc.add_paragraph(
    "The platform should not be described as fully smooth or fully ready today. It is, however, materially safer than a "
    "system that would silently score missing data. The next repairs should focus on the data spine: fresh official "
    "schedule metadata, provider identity mapping, complete park coverage, and a dashboard response that exposes weather "
    "already present in the database. Once those are complete, the correct indicator of success is a new, complete "
    "orchestration run—not merely a server that starts or a manually populated board."
)

doc.add_heading("Appendix A. Investigation Evidence", level=1)
add_table(doc,
          ["Evidence item", "Observed result"],
          [
              ["Today endpoint", "15 games; all weather labels = NOT FOUND; all roof labels = NOT FOUND; readiness = BLOCKED."],
              ["Weather table", "25 FantasyPros records across 15 games; 25 temperature values and 25 wind-speed values populated."],
              ["Weather refresh", "SUCCESS; 15 games found; 0 newly written; 0 failures. Existing FantasyPros fallback explains zero new writes."],
              ["Official games", "15 Aug. 25 games; 15 missing venue IDs; 15 missing start times."],
              ["Ballpark Pal latest run", "315 rows; 314 normalized; 1 rejected; PARTIAL; coverage gaps: game 0, hitter 35, pitcher 0, park 1."],
              ["Projected lineup / provider overlap", "270 projected FantasyPros hitters; 270 Ballpark Pal hitters; only 235 canonical MLB IDs overlap."],
              ["Runtime logs", "Database client termination events were observed around process lifecycle changes and deserve hardening investigation."],
          ],
          [2.1, 4.4])

doc.add_paragraph()
footer = doc.add_paragraph("This report separates confirmed observations from recommended remediation. It does not recommend odds, prices, EV, staking, or unsupported confidence claims.")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].italic = True
footer.runs[0].font.size = Pt(8)

doc.save(OUTPUT)
print(OUTPUT)